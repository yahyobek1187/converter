import os
import asyncio
from typing import List, Optional, Dict
import logging
from pathlib import Path

# Import conversion libraries
try:
    from pdf2docx import Converter as PDFConverter
    from docx import Document
    from PIL import Image
    from pydub import AudioSegment
    try:
        from moviepy.editor import VideoFileClip
        MOVIEPY_AVAILABLE = True
    except ImportError:
        VideoFileClip = None
        MOVIEPY_AVAILABLE = False
        logging.warning("MoviePy not available - video conversion disabled")
except ImportError as e:
    logging.error(f"Required library not found: {e}")
    raise

logger = logging.getLogger(__name__)

# Conversion matrix - defines which formats can be converted to which
def get_conversion_matrix():
    base_matrix = {
        '.pdf': ['docx', 'txt'],
        '.docx': ['pdf', 'txt'],
        '.txt': ['docx'],
        '.jpg': ['png', 'webp'],
        '.jpeg': ['png', 'webp'],
        '.png': ['jpg', 'webp'],
        '.webp': ['jpg', 'png'],
        '.mp3': ['wav', 'ogg'],
        '.wav': ['mp3', 'ogg'],
        '.ogg': ['mp3', 'wav'],
    }
    if MOVIEPY_AVAILABLE:
        base_matrix['.mp4'] = ['mp3']
    return base_matrix

CONVERSION_MATRIX = get_conversion_matrix()

def get_supported_conversions(file_extension: str) -> List[str]:
    """Get list of supported conversion formats for a given file extension."""
    return CONVERSION_MATRIX.get(file_extension.lower(), [])

def cleanup_temp_files(file_paths: List[str]) -> None:
    """Clean up temporary files."""
    for file_path in file_paths:
        try:
            if os.path.exists(file_path):
                os.remove(file_path)
                logger.info(f"Cleaned up temporary file: {file_path}")
        except Exception as e:
            logger.error(f"Error cleaning up file {file_path}: {e}")

class FileConverter:
    """File conversion utility class."""
    
    def __init__(self):
        self.temp_dir = "temp"
        os.makedirs(self.temp_dir, exist_ok=True)
    
    async def convert_file(self, input_path: str, source_format: str, target_format: str) -> Optional[str]:
        """
        Convert a file from source format to target format.
        Returns the path to the converted file or None if conversion failed.
        """
        try:
            # Generate output path
            base_name = os.path.splitext(os.path.basename(input_path))[0]
            output_path = os.path.join(self.temp_dir, f"{base_name}_converted.{target_format}")
            
            # Route to appropriate conversion method
            if source_format in ['pdf'] and target_format in ['docx', 'txt']:
                return await self._convert_pdf(input_path, output_path, target_format)
            elif source_format in ['docx'] and target_format in ['pdf', 'txt']:
                return await self._convert_docx(input_path, output_path, target_format)
            elif source_format in ['txt'] and target_format in ['docx']:
                return await self._convert_txt_to_docx(input_path, output_path)
            elif source_format in ['jpg', 'jpeg', 'png', 'webp'] and target_format in ['jpg', 'png', 'webp']:
                return await self._convert_image(input_path, output_path, target_format)
            elif source_format in ['mp3', 'wav', 'ogg'] and target_format in ['mp3', 'wav', 'ogg']:
                return await self._convert_audio(input_path, output_path, target_format)
            elif source_format == 'mp4' and target_format == 'mp3':
                return await self._convert_video_to_audio(input_path, output_path)
            else:
                logger.error(f"Unsupported conversion: {source_format} -> {target_format}")
                return None
                
        except Exception as e:
            logger.error(f"Conversion error: {e}")
            return None
    
    async def _convert_pdf(self, input_path: str, output_path: str, target_format: str) -> Optional[str]:
        """Convert PDF to DOCX or TXT."""
        try:
            if target_format == 'docx':
                # Convert PDF to DOCX
                cv = PDFConverter(input_path)
                cv.convert(output_path, start=0, end=None)
                cv.close()
            elif target_format == 'txt':
                # Convert PDF to TXT via intermediate DOCX
                temp_docx = output_path.replace('.txt', '_temp.docx')
                cv = PDFConverter(input_path)
                cv.convert(temp_docx, start=0, end=None)
                cv.close()
                
                # Extract text from DOCX
                doc = Document(temp_docx)
                text_content = '\n'.join([paragraph.text for paragraph in doc.paragraphs])
                
                with open(output_path, 'w', encoding='utf-8') as f:
                    f.write(text_content)
                
                # Clean up temp DOCX
                os.remove(temp_docx)
            
            return output_path if os.path.exists(output_path) else None
            
        except Exception as e:
            logger.error(f"PDF conversion error: {e}")
            return None
    
    async def _convert_docx(self, input_path: str, output_path: str, target_format: str) -> Optional[str]:
        """Convert DOCX to PDF or TXT."""
        try:
            if target_format == 'txt':
                # Extract text from DOCX
                doc = Document(input_path)
                text_content = '\n'.join([paragraph.text for paragraph in doc.paragraphs])
                
                with open(output_path, 'w', encoding='utf-8') as f:
                    f.write(text_content)
                    
            elif target_format == 'pdf':
                # For DOCX to PDF, we'll need to use a different approach
                # Since python-docx2pdf might not be available, we'll convert via text
                doc = Document(input_path)
                text_content = '\n'.join([paragraph.text for paragraph in doc.paragraphs])
                
                # Create a simple text-based PDF using reportlab if available
                try:
                    from reportlab.pdfgen import canvas
                    from reportlab.lib.pagesizes import letter
                    
                    c = canvas.Canvas(output_path, pagesize=letter)
                    width, height = letter
                    
                    # Split text into lines that fit the page
                    lines = text_content.split('\n')
                    y_position = height - 50
                    
                    for line in lines:
                        if y_position < 50:  # Start new page if needed
                            c.showPage()
                            y_position = height - 50
                        
                        # Split long lines
                        if len(line) > 80:
                            words = line.split(' ')
                            current_line = ''
                            for word in words:
                                if len(current_line + word) < 80:
                                    current_line += word + ' '
                                else:
                                    if current_line:
                                        c.drawString(50, y_position, current_line.strip())
                                        y_position -= 15
                                    current_line = word + ' '
                            if current_line:
                                c.drawString(50, y_position, current_line.strip())
                                y_position -= 15
                        else:
                            c.drawString(50, y_position, line)
                            y_position -= 15
                    
                    c.save()
                    
                except ImportError:
                    # Fallback: create a text file with PDF extension
                    with open(output_path, 'w', encoding='utf-8') as f:
                        f.write(text_content)
            
            return output_path if os.path.exists(output_path) else None
            
        except Exception as e:
            logger.error(f"DOCX conversion error: {e}")
            return None
    
    async def _convert_txt_to_docx(self, input_path: str, output_path: str) -> Optional[str]:
        """Convert TXT to DOCX."""
        try:
            # Read text file
            with open(input_path, 'r', encoding='utf-8') as f:
                text_content = f.read()
            
            # Create DOCX document
            doc = Document()
            
            # Split text into paragraphs and add to document
            paragraphs = text_content.split('\n\n')
            for paragraph in paragraphs:
                if paragraph.strip():
                    doc.add_paragraph(paragraph.strip())
            
            doc.save(output_path)
            return output_path if os.path.exists(output_path) else None
            
        except Exception as e:
            logger.error(f"TXT to DOCX conversion error: {e}")
            return None
    
    async def _convert_image(self, input_path: str, output_path: str, target_format: str) -> Optional[str]:
        """Convert between image formats."""
        try:
            with Image.open(input_path) as img:
                # Convert RGBA to RGB for JPEG
                if target_format.lower() in ['jpg', 'jpeg'] and img.mode in ('RGBA', 'LA', 'P'):
                    background = Image.new('RGB', img.size, (255, 255, 255))
                    if img.mode == 'P':
                        img = img.convert('RGBA')
                    background.paste(img, mask=img.split()[-1] if img.mode == 'RGBA' else None)
                    img = background
                
                # Save in target format
                save_format = 'JPEG' if target_format.lower() in ['jpg', 'jpeg'] else target_format.upper()
                img.save(output_path, format=save_format, quality=95 if save_format == 'JPEG' else None)
            
            return output_path if os.path.exists(output_path) else None
            
        except Exception as e:
            logger.error(f"Image conversion error: {e}")
            return None
    
    async def _convert_audio(self, input_path: str, output_path: str, target_format: str) -> Optional[str]:
        """Convert between audio formats."""
        try:
            audio = AudioSegment.from_file(input_path)
            
            # Export in target format
            if target_format == 'mp3':
                audio.export(output_path, format='mp3', bitrate='192k')
            elif target_format == 'wav':
                audio.export(output_path, format='wav')
            elif target_format == 'ogg':
                audio.export(output_path, format='ogg')
            
            return output_path if os.path.exists(output_path) else None
            
        except Exception as e:
            logger.error(f"Audio conversion error: {e}")
            return None
    
    async def _convert_video_to_audio(self, input_path: str, output_path: str) -> Optional[str]:
        """Convert video to audio (MP4 to MP3)."""
        try:
            if not MOVIEPY_AVAILABLE:
                logger.error("MoviePy not available - cannot convert video to audio")
                return None
                
            with VideoFileClip(input_path) as video:
                audio = video.audio
                if audio:
                    audio.write_audiofile(output_path, verbose=False, logger=None)
                    audio.close()
                else:
                    logger.error("Video file has no audio track")
                    return None
            
            return output_path if os.path.exists(output_path) else None
            
        except Exception as e:
            logger.error(f"Video to audio conversion error: {e}")
            return None
